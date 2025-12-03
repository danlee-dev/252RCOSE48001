"""
Multi-Format Document Parser (Production-Grade)
- PDF, HWP, DOCX, TXT, 이미지 파일 지원
- Vision API를 통한 OCR 및 구조화 추출
- 파일 형식별 최적화된 텍스트 추출

Supported Formats:
- PDF: pypdfium2 + Vision OCR fallback
- HWP: olefile + pyhwp (한글 문서)
- DOCX: python-docx (MS Word)
- TXT: Plain text (UTF-8, EUC-KR 자동 감지)
- Images: PNG, JPG, JPEG, GIF, WEBP, BMP, TIFF (Vision API)
"""

import os
import re
import io
import base64
import tempfile
import struct
import zlib
from pathlib import Path
from typing import Dict, Any, Optional, List, Tuple, Union
from dataclasses import dataclass, field
from enum import Enum
from abc import ABC, abstractmethod


class DocumentType(Enum):
    """지원 문서 형식"""
    PDF = "pdf"
    HWP = "hwp"
    HWPX = "hwpx"
    DOCX = "docx"
    DOC = "doc"
    TXT = "txt"
    RTF = "rtf"
    IMAGE = "image"
    UNKNOWN = "unknown"


@dataclass
class ParsedDocument:
    """파싱된 문서 결과"""
    text: str
    document_type: DocumentType
    page_count: int = 1
    metadata: Dict[str, Any] = field(default_factory=dict)
    tables: List[Dict[str, Any]] = field(default_factory=list)
    images: List[Dict[str, Any]] = field(default_factory=list)
    structured_content: Optional[str] = None
    confidence: float = 1.0
    warnings: List[str] = field(default_factory=list)

    @property
    def is_empty(self) -> bool:
        return not self.text or len(self.text.strip()) < 10

    @property
    def char_count(self) -> int:
        return len(self.text) if self.text else 0

    def to_dict(self) -> Dict[str, Any]:
        return {
            "text": self.text,
            "document_type": self.document_type.value,
            "page_count": self.page_count,
            "metadata": self.metadata,
            "tables": self.tables,
            "char_count": self.char_count,
            "confidence": self.confidence,
            "warnings": self.warnings
        }


class BaseParser(ABC):
    """문서 파서 기본 클래스"""

    @abstractmethod
    def parse(self, file_path: str) -> ParsedDocument:
        pass

    @abstractmethod
    def can_parse(self, file_path: str) -> bool:
        pass


class PDFParser(BaseParser):
    """PDF 문서 파서"""

    def can_parse(self, file_path: str) -> bool:
        return Path(file_path).suffix.lower() == ".pdf"

    def parse(self, file_path: str, use_ocr: bool = False) -> ParsedDocument:
        """
        PDF 파일 파싱

        1차: pdfplumber로 텍스트 추출
        2차: 텍스트 부족 시 pypdfium2 + Vision OCR
        """
        warnings = []
        text = ""
        page_count = 0
        tables = []

        # 1차: pdfplumber 시도
        try:
            import pdfplumber

            with pdfplumber.open(file_path) as pdf:
                page_count = len(pdf.pages)
                page_texts = []

                for i, page in enumerate(pdf.pages):
                    page_text = page.extract_text() or ""
                    page_texts.append(f"--- Page {i + 1} ---\n{page_text}")

                    # 표 추출
                    page_tables = page.extract_tables()
                    for table in page_tables:
                        if table:
                            tables.append({
                                "page": i + 1,
                                "data": table
                            })

                text = "\n\n".join(page_texts)

        except ImportError:
            warnings.append("pdfplumber not installed")
        except Exception as e:
            warnings.append(f"pdfplumber error: {str(e)}")

        # 텍스트가 부족하면 OCR 사용
        if len(text.strip()) < 100 or use_ocr:
            try:
                from app.ai.vision_parser import VisionParser

                parser = VisionParser()
                ocr_text = parser.extract_text_from_pdf(file_path)

                if len(ocr_text) > len(text):
                    text = ocr_text
                    warnings.append("Vision OCR used for text extraction")

            except Exception as e:
                warnings.append(f"Vision OCR error: {str(e)}")

        return ParsedDocument(
            text=text,
            document_type=DocumentType.PDF,
            page_count=page_count,
            tables=tables,
            warnings=warnings,
            metadata={"source": "pdfplumber" if "OCR" not in str(warnings) else "vision_ocr"}
        )


class HWPParser(BaseParser):
    """HWP 한글 문서 파서"""

    # HWP 파일 시그니처
    HWP_SIGNATURE = b'\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1'  # OLE compound document

    def can_parse(self, file_path: str) -> bool:
        suffix = Path(file_path).suffix.lower()
        return suffix in [".hwp", ".hwpx"]

    def parse(self, file_path: str) -> ParsedDocument:
        """
        HWP 파일 파싱

        1차: pyhwp 라이브러리 사용
        2차: olefile로 직접 파싱
        3차: hwp5txt 명령줄 도구 사용
        4차: Vision OCR fallback (이미지 변환)
        """
        suffix = Path(file_path).suffix.lower()
        warnings = []
        text = ""

        if suffix == ".hwpx":
            # HWPX는 ZIP 기반 (OOXML 유사)
            text, warnings = self._parse_hwpx(file_path)
        else:
            # HWP (OLE compound document)
            text, warnings = self._parse_hwp(file_path)

        # 텍스트가 부족하면 Vision OCR 시도
        if len(text.strip()) < 50:
            try:
                ocr_text = self._hwp_to_image_ocr(file_path)
                if len(ocr_text) > len(text):
                    text = ocr_text
                    warnings.append("Vision OCR used for HWP extraction")
            except Exception as e:
                warnings.append(f"HWP OCR fallback error: {str(e)}")

        return ParsedDocument(
            text=text,
            document_type=DocumentType.HWPX if suffix == ".hwpx" else DocumentType.HWP,
            warnings=warnings,
            metadata={"format": suffix}
        )

    def _parse_hwp(self, file_path: str) -> Tuple[str, List[str]]:
        """HWP 파일 파싱 (OLE compound document)"""
        warnings = []
        text = ""

        # 방법 1: pyhwp 라이브러리
        try:
            import hwp5.hwp5txt
            from io import StringIO
            import sys

            old_stdout = sys.stdout
            sys.stdout = mystdout = StringIO()

            try:
                hwp5.hwp5txt.main([file_path])
                text = mystdout.getvalue()
            finally:
                sys.stdout = old_stdout

            if text.strip():
                return text, warnings

        except ImportError:
            warnings.append("pyhwp not installed (pip install pyhwp)")
        except Exception as e:
            warnings.append(f"pyhwp error: {str(e)}")

        # 방법 2: olefile로 직접 파싱
        try:
            import olefile

            ole = olefile.OleFileIO(file_path)

            # HWP의 BodyText 스트림에서 텍스트 추출
            text_parts = []

            for stream_name in ole.listdir():
                stream_path = "/".join(stream_name)

                if "BodyText" in stream_path or "Section" in stream_path:
                    try:
                        data = ole.openstream(stream_name).read()
                        # HWP 압축 해제 및 텍스트 추출
                        extracted = self._extract_hwp_text_from_stream(data)
                        if extracted:
                            text_parts.append(extracted)
                    except Exception:
                        continue

            ole.close()

            if text_parts:
                text = "\n".join(text_parts)
                return text, warnings

        except ImportError:
            warnings.append("olefile not installed (pip install olefile)")
        except Exception as e:
            warnings.append(f"olefile error: {str(e)}")

        # 방법 3: hwp5txt 명령줄 도구
        try:
            import subprocess
            result = subprocess.run(
                ["hwp5txt", file_path],
                capture_output=True,
                text=True,
                timeout=30
            )
            if result.returncode == 0 and result.stdout.strip():
                text = result.stdout
                return text, warnings
        except FileNotFoundError:
            warnings.append("hwp5txt command not found")
        except Exception as e:
            warnings.append(f"hwp5txt error: {str(e)}")

        return text, warnings

    def _extract_hwp_text_from_stream(self, data: bytes) -> str:
        """HWP 스트림에서 텍스트 추출 (HWP 5.0 포맷)"""
        text_parts = []

        try:
            # 압축된 경우 zlib으로 해제
            if data[:2] == b'\x1f\x8b':  # gzip
                import gzip
                data = gzip.decompress(data)
            elif len(data) > 2:
                try:
                    data = zlib.decompress(data, -15)
                except:
                    pass

            # HWP 5.0 레코드 구조 파싱
            # 레코드 헤더: TagID(10bit) + Level(10bit) + Size(12bit) = 4bytes
            pos = 0
            while pos < len(data) - 4:
                try:
                    header = struct.unpack('<I', data[pos:pos+4])[0]
                    tag_id = header & 0x3FF
                    size = (header >> 20) & 0xFFF

                    # 확장 크기 처리
                    if size == 0xFFF:
                        if pos + 8 > len(data):
                            break
                        size = struct.unpack('<I', data[pos+4:pos+8])[0]
                        pos += 8
                    else:
                        pos += 4

                    if pos + size > len(data):
                        break

                    record_data = data[pos:pos+size]
                    pos += size

                    # HWPTAG_PARA_TEXT (67) - 텍스트 레코드
                    if tag_id == 67 and size >= 2:
                        # UTF-16LE로 텍스트 추출
                        try:
                            text = record_data.decode('utf-16le', errors='ignore')
                            # 제어 문자 필터링 (HWP 제어 코드)
                            cleaned = []
                            for char in text:
                                code = ord(char)
                                # 일반 텍스트만 추출 (제어 코드 제외)
                                if code >= 32 or char in '\n\r\t':
                                    if char not in '\x00\x01\x02\x03\x04\x05\x06\x07\x08\x0b\x0c\x0e\x0f':
                                        cleaned.append(char)
                                elif code == 10:  # 줄바꿈
                                    cleaned.append('\n')
                            if cleaned:
                                text_parts.append(''.join(cleaned))
                        except:
                            pass

                except struct.error:
                    break
                except Exception:
                    pos += 1
                    continue

            text = '\n'.join(filter(None, text_parts))
            # 불필요한 공백 정리
            text = re.sub(r'\n{3,}', '\n\n', text)
            text = re.sub(r'[ \t]+', ' ', text)
            return text.strip()

        except Exception:
            return ""

    def _parse_hwpx(self, file_path: str) -> Tuple[str, List[str]]:
        """HWPX 파일 파싱 (ZIP 기반)"""
        warnings = []
        text_parts = []

        try:
            import zipfile
            from xml.etree import ElementTree as ET

            with zipfile.ZipFile(file_path, 'r') as hwpx:
                for name in hwpx.namelist():
                    if 'section' in name.lower() and name.endswith('.xml'):
                        try:
                            with hwpx.open(name) as f:
                                content = f.read()
                                # XML에서 텍스트 추출
                                root = ET.fromstring(content)
                                text = self._extract_text_from_xml(root)
                                if text:
                                    text_parts.append(text)
                        except Exception:
                            continue

        except ImportError:
            warnings.append("zipfile error")
        except Exception as e:
            warnings.append(f"HWPX parsing error: {str(e)}")

        return "\n".join(text_parts), warnings

    def _extract_text_from_xml(self, element) -> str:
        """XML 요소에서 텍스트 재귀 추출"""
        texts = []
        if element.text:
            texts.append(element.text.strip())
        for child in element:
            texts.append(self._extract_text_from_xml(child))
        if element.tail:
            texts.append(element.tail.strip())
        return " ".join(filter(None, texts))

    def _hwp_to_image_ocr(self, file_path: str) -> str:
        """HWP를 이미지로 변환 후 OCR (LibreOffice 활용)"""
        try:
            import subprocess
            from app.ai.vision_parser import VisionParser

            # LibreOffice로 PDF 변환
            with tempfile.TemporaryDirectory() as tmpdir:
                result = subprocess.run([
                    "libreoffice",
                    "--headless",
                    "--convert-to", "pdf",
                    "--outdir", tmpdir,
                    file_path
                ], capture_output=True, timeout=60)

                if result.returncode == 0:
                    pdf_path = Path(tmpdir) / (Path(file_path).stem + ".pdf")
                    if pdf_path.exists():
                        parser = VisionParser()
                        return parser.extract_text_from_pdf(str(pdf_path))

        except Exception as e:
            pass

        return ""


class DOCXParser(BaseParser):
    """DOCX 문서 파서"""

    def can_parse(self, file_path: str) -> bool:
        suffix = Path(file_path).suffix.lower()
        return suffix in [".docx", ".doc"]

    def parse(self, file_path: str) -> ParsedDocument:
        """
        DOCX/DOC 파일 파싱

        DOCX: python-docx 사용
        DOC: antiword 또는 LibreOffice 변환
        """
        suffix = Path(file_path).suffix.lower()
        warnings = []
        text = ""
        tables = []

        if suffix == ".docx":
            text, tables, warnings = self._parse_docx(file_path)
        else:
            text, warnings = self._parse_doc(file_path)

        return ParsedDocument(
            text=text,
            document_type=DocumentType.DOCX if suffix == ".docx" else DocumentType.DOC,
            tables=tables,
            warnings=warnings,
            metadata={"format": suffix}
        )

    def _parse_docx(self, file_path: str) -> Tuple[str, List[Dict], List[str]]:
        """DOCX 파싱"""
        warnings = []
        text_parts = []
        tables = []

        try:
            from docx import Document

            doc = Document(file_path)

            # 문단 추출
            for para in doc.paragraphs:
                if para.text.strip():
                    # 스타일 정보 포함
                    style = para.style.name if para.style else "Normal"
                    if "Heading" in style:
                        text_parts.append(f"\n## {para.text}\n")
                    else:
                        text_parts.append(para.text)

            # 표 추출
            for i, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                if table_data:
                    tables.append({
                        "index": i,
                        "data": table_data
                    })

            text = "\n".join(text_parts)

        except ImportError:
            warnings.append("python-docx not installed (pip install python-docx)")
        except Exception as e:
            warnings.append(f"DOCX parsing error: {str(e)}")

        return text, tables, warnings

    def _parse_doc(self, file_path: str) -> Tuple[str, List[str]]:
        """DOC 파싱 (레거시 Word 형식)"""
        warnings = []
        text = ""

        # 방법 1: antiword
        try:
            import subprocess
            result = subprocess.run(
                ["antiword", file_path],
                capture_output=True,
                text=True,
                timeout=30
            )
            if result.returncode == 0 and result.stdout.strip():
                text = result.stdout
                return text, warnings
        except FileNotFoundError:
            warnings.append("antiword not found")
        except Exception as e:
            warnings.append(f"antiword error: {str(e)}")

        # 방법 2: LibreOffice 변환
        try:
            import subprocess

            with tempfile.TemporaryDirectory() as tmpdir:
                result = subprocess.run([
                    "libreoffice",
                    "--headless",
                    "--convert-to", "txt:Text",
                    "--outdir", tmpdir,
                    file_path
                ], capture_output=True, timeout=60)

                if result.returncode == 0:
                    txt_path = Path(tmpdir) / (Path(file_path).stem + ".txt")
                    if txt_path.exists():
                        text = txt_path.read_text(encoding='utf-8', errors='ignore')
                        return text, warnings

        except Exception as e:
            warnings.append(f"LibreOffice conversion error: {str(e)}")

        return text, warnings


class TXTParser(BaseParser):
    """텍스트 파일 파서"""

    # 지원 인코딩 목록 (우선순위 순)
    ENCODINGS = ['utf-8', 'utf-8-sig', 'euc-kr', 'cp949', 'latin-1', 'iso-8859-1']

    def can_parse(self, file_path: str) -> bool:
        suffix = Path(file_path).suffix.lower()
        return suffix in [".txt", ".text", ".md", ".markdown", ".rtf"]

    def parse(self, file_path: str) -> ParsedDocument:
        """텍스트 파일 파싱 (인코딩 자동 감지)"""
        warnings = []
        text = ""
        detected_encoding = None

        # 방법 1: chardet으로 인코딩 감지
        try:
            import chardet
            with open(file_path, 'rb') as f:
                raw_data = f.read()
            detected = chardet.detect(raw_data)
            if detected['confidence'] > 0.7:
                detected_encoding = detected['encoding']
                text = raw_data.decode(detected_encoding, errors='ignore')
        except ImportError:
            warnings.append("chardet not installed, using fallback encoding detection")
        except Exception as e:
            warnings.append(f"chardet error: {str(e)}")

        # 방법 2: 수동 인코딩 시도
        if not text:
            for encoding in self.ENCODINGS:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        text = f.read()
                    detected_encoding = encoding
                    break
                except (UnicodeDecodeError, LookupError):
                    continue

        # RTF 처리
        if Path(file_path).suffix.lower() == ".rtf":
            text = self._strip_rtf(text)

        return ParsedDocument(
            text=text,
            document_type=DocumentType.TXT,
            warnings=warnings,
            metadata={"encoding": detected_encoding or "unknown"}
        )

    def _strip_rtf(self, rtf_text: str) -> str:
        """RTF 마크업 제거"""
        # 간단한 RTF 스트리퍼
        import re

        # RTF 제어 단어 제거
        text = re.sub(r'\\[a-z]+\d*\s?', '', rtf_text)
        # 중괄호 제거
        text = re.sub(r'[{}]', '', text)
        # 특수 문자 변환
        text = text.replace('\\par', '\n')
        text = text.replace('\\tab', '\t')

        return text.strip()


class ImageParser(BaseParser):
    """이미지 파일 파서 (Vision API OCR)"""

    # 지원 이미지 확장자
    IMAGE_EXTENSIONS = {'.png', '.jpg', '.jpeg', '.gif', '.webp', '.bmp', '.tiff', '.tif'}

    def can_parse(self, file_path: str) -> bool:
        suffix = Path(file_path).suffix.lower()
        return suffix in self.IMAGE_EXTENSIONS

    def parse(self, file_path: str, structured: bool = True) -> ParsedDocument:
        """
        이미지 파일 파싱 (Vision API 사용)

        Args:
            file_path: 이미지 파일 경로
            structured: True면 구조화된 마크다운 추출, False면 텍스트만
        """
        warnings = []
        text = ""
        tables = []
        structured_content = None

        try:
            from app.ai.vision_parser import VisionParser

            parser = VisionParser()

            if structured:
                result = parser.parse_image(file_path, extract_tables=True)
                text = result.raw_text
                structured_content = result.structured_markdown
                tables = result.tables
            else:
                text = parser.extract_text_only(file_path)

        except Exception as e:
            warnings.append(f"Vision OCR error: {str(e)}")

        return ParsedDocument(
            text=text,
            document_type=DocumentType.IMAGE,
            tables=tables,
            structured_content=structured_content,
            warnings=warnings,
            metadata={"format": Path(file_path).suffix.lower()}
        )


class MultiFormatDocumentParser:
    """
    멀티포맷 문서 파서 (Production-Grade)

    지원 형식:
    - PDF: pdfplumber + Vision OCR fallback
    - HWP/HWPX: pyhwp, olefile, hwp5txt
    - DOCX/DOC: python-docx, antiword, LibreOffice
    - TXT/MD/RTF: 자동 인코딩 감지
    - Images: Vision API OCR

    사용법:
        parser = MultiFormatDocumentParser()
        result = parser.parse("contract.hwp")
        print(result.text)
        print(result.tables)
    """

    # 지원 MIME 타입 매핑
    MIME_TYPE_MAP = {
        "application/pdf": DocumentType.PDF,
        "application/x-hwp": DocumentType.HWP,
        "application/haansofthwp": DocumentType.HWP,
        "application/vnd.hancom.hwp": DocumentType.HWP,
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": DocumentType.DOCX,
        "application/msword": DocumentType.DOC,
        "text/plain": DocumentType.TXT,
        "text/markdown": DocumentType.TXT,
        "application/rtf": DocumentType.TXT,
        "image/png": DocumentType.IMAGE,
        "image/jpeg": DocumentType.IMAGE,
        "image/gif": DocumentType.IMAGE,
        "image/webp": DocumentType.IMAGE,
        "image/bmp": DocumentType.IMAGE,
        "image/tiff": DocumentType.IMAGE,
    }

    # 확장자 -> DocumentType 매핑
    EXTENSION_MAP = {
        ".pdf": DocumentType.PDF,
        ".hwp": DocumentType.HWP,
        ".hwpx": DocumentType.HWPX,
        ".docx": DocumentType.DOCX,
        ".doc": DocumentType.DOC,
        ".txt": DocumentType.TXT,
        ".text": DocumentType.TXT,
        ".md": DocumentType.TXT,
        ".markdown": DocumentType.TXT,
        ".rtf": DocumentType.TXT,
        ".png": DocumentType.IMAGE,
        ".jpg": DocumentType.IMAGE,
        ".jpeg": DocumentType.IMAGE,
        ".gif": DocumentType.IMAGE,
        ".webp": DocumentType.IMAGE,
        ".bmp": DocumentType.IMAGE,
        ".tiff": DocumentType.IMAGE,
        ".tif": DocumentType.IMAGE,
    }

    def __init__(self):
        self.parsers = {
            DocumentType.PDF: PDFParser(),
            DocumentType.HWP: HWPParser(),
            DocumentType.HWPX: HWPParser(),
            DocumentType.DOCX: DOCXParser(),
            DocumentType.DOC: DOCXParser(),
            DocumentType.TXT: TXTParser(),
            DocumentType.IMAGE: ImageParser(),
        }

    def detect_type(self, file_path: str, mime_type: str = None) -> DocumentType:
        """파일 형식 감지"""

        # MIME 타입으로 감지
        if mime_type and mime_type in self.MIME_TYPE_MAP:
            return self.MIME_TYPE_MAP[mime_type]

        # 확장자로 감지
        suffix = Path(file_path).suffix.lower()
        if suffix in self.EXTENSION_MAP:
            return self.EXTENSION_MAP[suffix]

        # 파일 시그니처로 감지
        return self._detect_by_signature(file_path)

    def _detect_by_signature(self, file_path: str) -> DocumentType:
        """파일 시그니처로 형식 감지"""
        try:
            with open(file_path, 'rb') as f:
                header = f.read(8)

            # PDF
            if header[:4] == b'%PDF':
                return DocumentType.PDF

            # OLE compound document (HWP, DOC)
            if header == b'\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1':
                # HWP인지 DOC인지 구분
                try:
                    import olefile
                    ole = olefile.OleFileIO(file_path)
                    if 'FileHeader' in ole.listdir() or 'HwpSummaryInformation' in str(ole.listdir()):
                        ole.close()
                        return DocumentType.HWP
                    ole.close()
                    return DocumentType.DOC
                except:
                    return DocumentType.DOC

            # ZIP (DOCX, HWPX)
            if header[:4] == b'PK\x03\x04':
                try:
                    import zipfile
                    with zipfile.ZipFile(file_path, 'r') as z:
                        names = z.namelist()
                        if any('word/' in n for n in names):
                            return DocumentType.DOCX
                        if any('Contents/' in n for n in names):
                            return DocumentType.HWPX
                except:
                    pass

            # PNG
            if header[:8] == b'\x89PNG\r\n\x1a\n':
                return DocumentType.IMAGE

            # JPEG
            if header[:2] == b'\xff\xd8':
                return DocumentType.IMAGE

            # GIF
            if header[:6] in (b'GIF87a', b'GIF89a'):
                return DocumentType.IMAGE

        except Exception:
            pass

        return DocumentType.UNKNOWN

    def parse(
        self,
        file_path: str,
        mime_type: str = None,
        use_ocr_fallback: bool = True
    ) -> ParsedDocument:
        """
        문서 파싱

        Args:
            file_path: 파일 경로
            mime_type: MIME 타입 (옵션)
            use_ocr_fallback: 텍스트 부족 시 OCR 사용

        Returns:
            ParsedDocument: 파싱 결과
        """
        doc_type = self.detect_type(file_path, mime_type)

        if doc_type == DocumentType.UNKNOWN:
            return ParsedDocument(
                text="",
                document_type=DocumentType.UNKNOWN,
                warnings=["Unknown document type"]
            )

        parser = self.parsers.get(doc_type)
        if not parser:
            return ParsedDocument(
                text="",
                document_type=doc_type,
                warnings=[f"No parser available for {doc_type.value}"]
            )

        result = parser.parse(file_path)

        # OCR fallback
        if use_ocr_fallback and result.is_empty and doc_type not in [DocumentType.IMAGE]:
            try:
                from app.ai.vision_parser import VisionParser

                # 문서를 PDF로 변환 후 OCR
                if doc_type in [DocumentType.HWP, DocumentType.HWPX, DocumentType.DOC, DocumentType.DOCX]:
                    ocr_text = self._convert_and_ocr(file_path)
                    if len(ocr_text) > len(result.text):
                        result.text = ocr_text
                        result.warnings.append("OCR fallback used after document conversion")

            except Exception as e:
                result.warnings.append(f"OCR fallback error: {str(e)}")

        return result

    def _convert_and_ocr(self, file_path: str) -> str:
        """문서를 PDF로 변환 후 OCR"""
        try:
            import subprocess
            from app.ai.vision_parser import VisionParser

            with tempfile.TemporaryDirectory() as tmpdir:
                result = subprocess.run([
                    "libreoffice",
                    "--headless",
                    "--convert-to", "pdf",
                    "--outdir", tmpdir,
                    file_path
                ], capture_output=True, timeout=60)

                if result.returncode == 0:
                    pdf_path = Path(tmpdir) / (Path(file_path).stem + ".pdf")
                    if pdf_path.exists():
                        parser = VisionParser()
                        return parser.extract_text_from_pdf(str(pdf_path))

        except Exception:
            pass

        return ""

    def get_supported_extensions(self) -> List[str]:
        """지원 확장자 목록 반환"""
        return list(self.EXTENSION_MAP.keys())

    def get_supported_mime_types(self) -> List[str]:
        """지원 MIME 타입 목록 반환"""
        return list(self.MIME_TYPE_MAP.keys())

    def is_supported(self, file_path: str = None, mime_type: str = None) -> bool:
        """파일 지원 여부 확인"""
        if mime_type and mime_type in self.MIME_TYPE_MAP:
            return True
        if file_path:
            suffix = Path(file_path).suffix.lower()
            return suffix in self.EXTENSION_MAP
        return False


# 편의 함수
def parse_document(file_path: str, mime_type: str = None) -> ParsedDocument:
    """간편 문서 파싱"""
    parser = MultiFormatDocumentParser()
    return parser.parse(file_path, mime_type)


def extract_text(file_path: str) -> str:
    """간편 텍스트 추출"""
    result = parse_document(file_path)
    return result.text


def get_supported_formats() -> Dict[str, List[str]]:
    """지원 형식 정보 반환"""
    parser = MultiFormatDocumentParser()
    return {
        "extensions": parser.get_supported_extensions(),
        "mime_types": parser.get_supported_mime_types()
    }
